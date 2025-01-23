import os
import re
import uuid
from docx import Document
from docx.shared import Inches
from flask import current_app, url_for
from werkzeug.utils import secure_filename
from wtforms.validators import ValidationError
import fitz
import pypandoc
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml


def save_file(images, pathname="UPLOAD_FOLDER"):
    save_file = ""
    if images:
        if len(images) > 0 and isinstance(images, (list, tuple)):
            for index, image_data in enumerate(images):
                single_file = save_single_file(image_data, pathname)
                save_file += single_file
                if index != len(images) - 1:
                    save_file += ","
        else:
            save_file = save_single_file(images, pathname)
    return save_file


def delete_file(pre_image, basepath="UPLOAD_FOLDER"):
    pre_path = current_app.config[basepath]
    if pre_image is not None and len(pre_image) > 0:
        if "," in pre_image:
            pre_images = pre_image.split(",")
            for img in pre_images:
                file_del = pre_path + "/" + img
                if os.path.exists(file_del):
                    os.remove(file_del)
        else:
            file_del = pre_path + "/" + pre_image
            if os.path.exists(file_del):
                os.remove(file_del)


def save_single_file(file, basepath="UPLOAD_FOLDER"):
    file_name = secure_filename(file.filename)
    if file_name != "":
        file_ext = os.path.splitext(file_name)[1]
        filename = str(uuid.uuid4()) + file_ext
        file.save(os.path.join(current_app.config[basepath], filename))
        return filename


def validate_file(form, field):
    if not field.data:
        raise ValidationError("文件为空.")
    
    if isinstance(field.data, (list, tuple)) and len(field.data) > 0:
        for data in field.data:
            mime_type = data.content_type
            allowed_mime_types = ["image/jpeg", "image/jpg", "image/png", "image/gif"]
            if mime_type not in allowed_mime_types:
                raise ValidationError("仅支持格式JPEG、PNG、GIF")
    else:
        data = field.data
        mime_type = data.content_type
        allowed_mime_types = ["image/jpeg", "image/jpg", "image/png", "image/gif"]
        if mime_type not in allowed_mime_types:
            raise ValidationError("仅支持格式JPEG、PNG、GIF")


def add_image_watermark_docx(source_file, watermark_file, target_text="签名"):
    doc = Document(os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], source_file))
    image_path = os.path.join(current_app.config["UPLOAD_FOLDER"], watermark_file)
    has_key_word = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if target_text in run.text:
                has_key_word = True
                run.text = run.text.replace(target_text, f"{target_text}\n")
                paragraph.add_run("\t")
                run_cell = paragraph.add_run()
                run_cell.add_picture(image_path, width=Inches(0.9), height=Inches(0.8))
    if not has_key_word:
        total_paragraphs = len(doc.paragraphs)
        paragraph_index = int(total_paragraphs * 0.75)
        paragraph = doc.paragraphs[paragraph_index]
        if len(paragraph.runs) > 0:
            run_index = len(paragraph.runs) - 1
        else:
            paragraph.add_run("\t")
            run_index = 0
        run = paragraph.runs[run_index]
        run.add_picture(image_path, width=Inches(1), height=Inches(1))  # 设置图片宽度

    out_file = f"{source_file}-sign.docx"
    doc.save(os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], out_file))
    return url_for("static", filename=f"docs/{out_file}", _external=True)


def convert_word_to_html(filename):
    try:
        input_file = os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], filename)
        # 使用 pypandoc 将 Word 转换为 HTML，并嵌入图片
        html_content = pypandoc.convert_file(
            input_file,
            "html",
            extra_args=["--wrap=none", "--self-contained", "--preserve-tabs"],
        )
        # 将空格替换为 &nbsp;，将空行替换为 <br>
        # html_content = html_content.replace(' ', '&nbsp;')  # 替换空格
        # html_content = html_content.replace('\n', '<br>')   # 替换空行
        # html_content = html_content.replace('\t', '&nbsp;&nbsp;&nbsp;&nbsp;')   # 替换空行
        html_content = re.sub(
            r"<style.*?>.*?</style>", "", html_content, flags=re.DOTALL
        )
        html_content = re.sub(
            r"<title.*?>.*?</title>", "", html_content, flags=re.DOTALL
        )
        html_content = re.sub(r"<meta.*?>", "", html_content, flags=re.DOTALL)

        html_content = html_content.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
    except Exception as e:
        print(f"转换失败: {e}")
        return ""
    return f"<div class='docx-contain'>{html_content}<div>"


def convert_word_to_pdf(filename):
    outfile_name = f"{filename}-wordparse.pdf"
    input_file = os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], filename)
    output_file = os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], outfile_name)
    try:
        # 使用 pypandoc 将 Word 转换为 PDF
        pypandoc.convert_file(
            input_file,
            "pdf",
            outputfile=output_file,
            extra_args=["--pdf-engine=xelatex"],
        )
        return url_for("static", filename=f"docs/{outfile_name}", _external=True)
    except Exception as e:
        print(f"转换失败: {e}")


def add_image_watermark_pdf(
    input_pdf, watermark_file, target_text="签名", alpha=0.5, scale=0.5
):
    watermark_path = os.path.join(current_app.config["UPLOAD_FOLDER"], watermark_file)
    watermark = fitz.open(watermark_path)
    watermark_page = watermark.load_page(0)
    watermark_rect = watermark_page.rect

    pdf_document = fitz.open(
        os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], input_pdf)
    )
    page_num = len(pdf_document) - 1
    page = pdf_document.load_page(page_num)
    page_rect = page.rect

    text_instances = page.search_for(target_text)
    if not text_instances:
        # 未找到 根据距离定位
        watermark_width = watermark_rect.width * scale
        watermark_height = watermark_rect.height * scale
        x = (page_rect.width - watermark_width) / 4 * 1
        y = (page_rect.height - watermark_height) / 4 * 3
        position = fitz.Rect(x, y, x + watermark_width, y + watermark_height)
        page.insert_image(position, filename=watermark_path, alpha=alpha)
    else:
        # 找到 根据文本定位
        text_rect = text_instances[-1]
        watermark_width = watermark_rect.width * scale
        watermark_height = watermark_rect.height * scale
        next_line_y = text_rect.y1
        x = text_rect.x0 + watermark_width / 4
        y = next_line_y + text_rect.height / 2
        position = fitz.Rect(x, y, x + watermark_width, y + watermark_height)
        page.insert_image(position, filename=watermark_path, alpha=alpha)

    output_pdf = f"{input_pdf}-sign.pdf"
    pdf_document.save(
        os.path.join(current_app.config["UPLOAD_FOLDER_DOCS"], output_pdf)
    )
    pdf_document.close()
    watermark.close()
    return url_for("static", filename=f"docs/{output_pdf}", _external=True)
